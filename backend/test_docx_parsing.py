"""Test DOCX file parsing"""
import os
from docx import Document

# Find the uploaded file
uploads_dir = "uploads"
if os.path.exists(uploads_dir):
    files = [f for f in os.listdir(uploads_dir) if f.endswith('.doc') or f.endswith('.docx')]
    print(f"Found files: {files}")
    
    for file in files:
        file_path = os.path.join(uploads_dir, file)
        print(f"\n{'='*60}")
        print(f"Testing: {file}")
        print(f"{'='*60}")
        
        try:
            doc = Document(file_path)
            
            # Extract from paragraphs
            print("\n--- Paragraphs ---")
            for i, para in enumerate(doc.paragraphs[:10]):  # First 10 paragraphs
                if para.text.strip():
                    print(f"{i+1}: {para.text[:100]}")
            
            # Extract from tables
            print(f"\n--- Tables ({len(doc.tables)} found) ---")
            for i, table in enumerate(doc.tables):
                print(f"\nTable {i+1}:")
                for j, row in enumerate(table.rows[:3]):  # First 3 rows
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        print(f"  Row {j+1}: {row_text[:100]}")
            
            # Get all text
            all_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    all_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            all_text.append(cell.text)
            
            full_text = '\n'.join(all_text)
            print(f"\n--- Summary ---")
            print(f"Total characters extracted: {len(full_text)}")
            print(f"Total words: {len(full_text.split())}")
            print(f"\nFirst 500 characters:\n{full_text[:500]}")
            
        except Exception as e:
            print(f"ERROR: {e}")
            import traceback
            traceback.print_exc()
else:
    print(f"Uploads directory '{uploads_dir}' not found")
