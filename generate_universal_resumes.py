"""
Enhanced script to generate all test resumes in DOCX and PDF formats
Covers diverse roles: Tech, Healthcare, Finance, Engineering, Marketing
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: reportlab not installed. PDF generation skipped.")


def read_txt_resume(filename):
    """Read text resume file"""
    filepath = os.path.join('test-resumes', filename)
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()


def create_docx_from_txt(txt_content, output_filename):
    """Create DOCX resume from text content"""
    doc = Document()
    
    lines = txt_content.split('\n')
    current_section = None
    
    for line in lines:
        line = line.strip()
        
        if not line:
            doc.add_paragraph()  # Empty line
            continue
        
        # Detect headings (ALL CAPS or ends with colon)
        if line.isupper() or (len(line) < 50 and ':' not in line and not line[0].isdigit()):
            heading = doc.add_heading(line, level=2)
            current_section = line
        elif line.endswith(':') and len(line) < 60:
            para = doc.add_paragraph(line)
            para.runs[0].bold = True
        else:
            # Regular paragraph
            if line.startswith('•') or line.startswith('-'):
                doc.add_paragraph(line[1:].strip(), style='List Bullet')
            elif line.startswith(('Email:', 'Phone:', 'Location:', 'Address:', 'License:')):
                para = doc.add_paragraph(line)
                para.runs[0].font.size = Pt(10)
            else:
                doc.add_paragraph(line)
    
    output_path = os.path.join('test-resumes', output_filename)
    doc.save(output_path)
    print(f"✅ Created: {output_filename}")
    return output_path


def create_pdf_from_txt(txt_content, output_filename):
    """Create PDF resume from text content"""
    if not PDF_AVAILABLE:
        print(f"⚠️  Skipped PDF: {output_filename} (reportlab not installed)")
        return None
    
    output_path = os.path.join('test-resumes', output_filename)
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    
    styles = getSampleStyleSheet()
    story = []
    
    # Custom styles
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading1'],
        fontSize=14,
        textColor='#2C3E50',
        spaceAfter=12,
        spaceBefore=12
    )
    
    normal_style = styles['Normal']
    normal_style.fontSize = 10
    normal_style.leading = 14
    
    lines = txt_content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            story.append(Spacer(1, 0.1*inch))
            continue
        
        # Detect headings
        if line.isupper() or (len(line) < 50 and ':' not in line):
            story.append(Paragraph(f"<b>{line}</b>", heading_style))
        else:
            # Escape special characters for XML
            line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(line, normal_style))
            story.append(Spacer(1, 0.05*inch))
    
    doc.build(story)
    print(f"✅ Created: {output_filename}")
    return output_path


def main():
    """Generate all resume formats"""
    
    print("=" * 60)
    print("Generating Universal Test Resumes")
    print("=" * 60)
    
    # List of all TXT resumes
    resumes = [
        'emily_rodriguez_data_scientist.txt',
        'david_thompson_marketing_manager.txt',
        'james_wilson_mechanical_engineer.txt',
        'sarah_johnson_nurse.txt',
        'robert_chen_financial_analyst.txt'
    ]
    
    total_created = 0
    
    for txt_file in resumes:
        print(f"\n📄 Processing: {txt_file}")
        
        try:
            # Read text content
            txt_content = read_txt_resume(txt_file)
            
            # Generate DOCX
            base_name = txt_file.replace('.txt', '')
            docx_file = f"{base_name}.docx"
            create_docx_from_txt(txt_content, docx_file)
            total_created += 1
            
            # Generate PDF
            pdf_file = f"{base_name}.pdf"
            create_pdf_from_txt(txt_content, pdf_file)
            total_created += 1
            
        except Exception as e:
            print(f"❌ Error processing {txt_file}: {str(e)}")
    
    print("\n" + "=" * 60)
    print(f"✅ Generation Complete!")
    print(f"Total files created: {total_created}")
    print(f"Test resumes directory: test-resumes/")
    print("=" * 60)


if __name__ == "__main__":
    main()
