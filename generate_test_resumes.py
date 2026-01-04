"""
Generate test resumes in different formats (DOCX and PDF)
This ensures we have proper test files for production testing
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Try to import reportlab for PDF generation
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: reportlab not installed. PDF generation skipped.")
    print("Install: pip install reportlab")


def create_docx_resume_fullstack():
    """Create Sarah Mitchell's full-stack developer resume in DOCX format"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Sarah Mitchell', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Senior Full-Stack Developer')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    
    # Contact Info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run('Email: sarah.mitchell@techmail.com | Phone: +1 (555) 987-6543 | San Francisco, CA')
    contact.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()  # Spacing
    
    # Professional Summary
    doc.add_heading('PROFESSIONAL SUMMARY', 1)
    doc.add_paragraph(
        'Highly skilled Full-Stack Developer with 5+ years of experience building scalable '
        'web applications and microservices. Expert in modern JavaScript frameworks, cloud '
        'architecture, and agile methodologies. Proven track record of leading development '
        'teams and delivering high-quality software solutions.'
    )
    
    # Technical Skills
    doc.add_heading('TECHNICAL SKILLS', 1)
    
    skills_table = doc.add_table(rows=7, cols=2)
    skills_table.style = 'Light Grid Accent 1'
    
    skills_data = [
        ('Languages:', 'JavaScript, TypeScript, Python, Java, SQL, HTML5, CSS3'),
        ('Frontend:', 'React, Next.js, Vue.js, Angular, Redux, Tailwind CSS, Material-UI'),
        ('Backend:', 'Node.js, Express, Django, Flask, Spring Boot, FastAPI'),
        ('Databases:', 'PostgreSQL, MongoDB, MySQL, Redis, DynamoDB'),
        ('Cloud & DevOps:', 'AWS, Docker, Kubernetes, Jenkins, CI/CD, Terraform'),
        ('Testing:', 'Jest, PyTest, Selenium, Cypress, Unit Testing, Integration Testing'),
        ('Tools:', 'Git, JIRA, Postman, VS Code, Figma, Slack')
    ]
    
    for idx, (category, skills) in enumerate(skills_data):
        skills_table.rows[idx].cells[0].text = category
        skills_table.rows[idx].cells[1].text = skills
    
    # Professional Experience
    doc.add_heading('PROFESSIONAL EXPERIENCE', 1)
    
    # Job 1
    job1_title = doc.add_paragraph()
    job1_title.add_run('Senior Full-Stack Developer | TechCorp Solutions\n').bold = True
    job1_title.add_run('San Francisco, CA | January 2022 - Present').italic = True
    
    responsibilities = [
        'Led development team of 8 engineers building enterprise SaaS platform serving 100K+ users',
        'Architected and implemented microservices infrastructure using Node.js, PostgreSQL, and AWS',
        'Reduced application load time by 60% through performance optimization and caching strategies',
        'Implemented CI/CD pipelines using Jenkins and Docker, reducing deployment time by 75%',
        'Mentored junior developers and conducted code reviews to maintain code quality standards',
        'Collaborated with product managers and designers using Agile/Scrum methodologies'
    ]
    
    for resp in responsibilities:
        doc.add_paragraph(resp, style='List Bullet')
    
    tech = doc.add_paragraph()
    tech.add_run('Technologies: ').bold = True
    tech.add_run('React, TypeScript, Node.js, PostgreSQL, AWS, Docker, Kubernetes')
    
    doc.add_paragraph()  # Spacing
    
    # Job 2
    job2_title = doc.add_paragraph()
    job2_title.add_run('Full-Stack Developer | StartupHub Inc.\n').bold = True
    job2_title.add_run('Palo Alto, CA | June 2019 - December 2021').italic = True
    
    responsibilities2 = [
        'Developed RESTful APIs and GraphQL endpoints serving mobile and web applications',
        'Built responsive frontend interfaces using React and Redux with excellent user experience',
        'Integrated third-party services including Stripe payments, SendGrid emails, and Auth0',
        'Wrote comprehensive unit tests and integration tests achieving 90% code coverage',
        'Participated in daily standups, sprint planning, and retrospectives'
    ]
    
    for resp in responsibilities2:
        doc.add_paragraph(resp, style='List Bullet')
    
    tech2 = doc.add_paragraph()
    tech2.add_run('Technologies: ').bold = True
    tech2.add_run('React, Node.js, MongoDB, AWS Lambda, Jest, Cypress')
    
    # Education
    doc.add_heading('EDUCATION', 1)
    edu = doc.add_paragraph()
    edu.add_run('Bachelor of Science in Computer Science\n').bold = True
    edu.add_run('University of California, Berkeley | 2014 - 2018\n')
    edu.add_run('GPA: 3.7/4.0')
    
    # Certifications
    doc.add_heading('CERTIFICATIONS', 1)
    certs = [
        'AWS Certified Solutions Architect - Associate (2023)',
        'MongoDB Certified Developer Associate (2022)',
        'Certified Kubernetes Application Developer (CKAD) (2021)',
        'Google Cloud Professional Cloud Architect (2023)'
    ]
    for cert in certs:
        doc.add_paragraph(cert, style='List Bullet')
    
    # Save
    output_path = 'test-resumes/sarah_mitchell_fullstack.docx'
    doc.save(output_path)
    print(f"✅ Created: {output_path}")
    return output_path


def create_docx_resume_devops():
    """Create Michael Chen's DevOps resume in DOCX format"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Michael Chen', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('DevOps Engineer & Cloud Architect')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    
    # Contact Info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run('michael.chen@cloudops.io | (555) 123-8899 | Seattle, WA')
    contact.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()  # Spacing
    
    # Professional Summary
    doc.add_heading('PROFESSIONAL SUMMARY', 1)
    doc.add_paragraph(
        'Certified DevOps Engineer with 6+ years experience designing and implementing cloud '
        'infrastructure, automation pipelines, and containerized applications. Expert in AWS, '
        'Kubernetes, Terraform, and CI/CD best practices. Strong background in system '
        'administration, security hardening, and cost optimization.'
    )
    
    # Core Competencies
    doc.add_heading('CORE COMPETENCIES', 1)
    
    skills_table = doc.add_table(rows=8, cols=2)
    skills_table.style = 'Light Grid Accent 1'
    
    skills_data = [
        ('Cloud Platforms:', 'AWS, Azure, Google Cloud Platform (GCP)'),
        ('Containers:', 'Docker, Kubernetes, ECS, EKS, Docker Swarm'),
        ('Infrastructure as Code:', 'Terraform, CloudFormation, Ansible, Pulumi'),
        ('CI/CD Tools:', 'Jenkins, GitLab CI, GitHub Actions, CircleCI, ArgoCD'),
        ('Monitoring:', 'Prometheus, Grafana, ELK Stack, Datadog, CloudWatch'),
        ('Scripting:', 'Bash, Python, PowerShell, Go'),
        ('Version Control:', 'Git, GitHub, GitLab, Bitbucket'),
        ('Security:', 'IAM, Security Groups, Secrets Management, SSL/TLS, Vault')
    ]
    
    for idx, (category, skills) in enumerate(skills_data):
        skills_table.rows[idx].cells[0].text = category
        skills_table.rows[idx].cells[1].text = skills
    
    # Work Experience
    doc.add_heading('WORK EXPERIENCE', 1)
    
    # Job 1
    job1_title = doc.add_paragraph()
    job1_title.add_run('Senior DevOps Engineer | CloudScale Technologies\n').bold = True
    job1_title.add_run('Seattle, WA | March 2021 - Present').italic = True
    
    responsibilities = [
        'Designed and deployed multi-region AWS infrastructure supporting 500K daily active users',
        'Implemented Infrastructure as Code using Terraform, reducing provisioning time from hours to minutes',
        'Built automated CI/CD pipelines with Jenkins and GitHub Actions, achieving 50+ deployments per week',
        'Migrated monolithic applications to microservices architecture using Kubernetes and Docker',
        'Set up comprehensive monitoring and alerting using Prometheus, Grafana, and PagerDuty',
        'Reduced cloud infrastructure costs by 40% through resource optimization and auto-scaling policies',
        'Led security audits and implemented compliance measures (SOC2, GDPR, HIPAA)'
    ]
    
    for resp in responsibilities:
        doc.add_paragraph(resp, style='List Bullet')
    
    tech = doc.add_paragraph()
    tech.add_run('Technologies: ').bold = True
    tech.add_run('AWS, Kubernetes, Docker, Terraform, Jenkins, Python, Bash')
    
    # Education
    doc.add_heading('EDUCATION', 1)
    edu = doc.add_paragraph()
    edu.add_run('Bachelor of Science in Information Technology\n').bold = True
    edu.add_run('University of Washington | 2013 - 2017\n')
    edu.add_run('Major: Network Administration and Security | GPA: 3.6/4.0')
    
    # Certifications
    doc.add_heading('CERTIFICATIONS', 1)
    certs = [
        'AWS Certified Solutions Architect - Professional (2023)',
        'Certified Kubernetes Administrator (CKA) - 2022',
        'AWS Certified DevOps Engineer - Professional (2021)',
        'HashiCorp Certified: Terraform Associate (2022)',
        'AWS Certified Security - Specialty (2023)'
    ]
    for cert in certs:
        doc.add_paragraph(cert, style='List Bullet')
    
    # Save
    output_path = 'test-resumes/michael_chen_devops.docx'
    doc.save(output_path)
    print(f"✅ Created: {output_path}")
    return output_path


def create_pdf_resume_fullstack():
    """Create Sarah Mitchell's resume in PDF format"""
    if not PDF_AVAILABLE:
        print("⚠️  Skipping PDF creation - reportlab not installed")
        return None
    
    output_path = 'test-resumes/sarah_mitchell_fullstack.pdf'
    doc = SimpleDocTemplate(output_path, pagesize=letter,
                           rightMargin=0.75*inch, leftMargin=0.75*inch,
                           topMargin=0.75*inch, bottomMargin=0.75*inch)
    
    # Container for the 'Flowable' objects
    elements = []
    
    # Define styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=RGBColor(0, 0, 0),
        spaceAfter=6,
        alignment=TA_CENTER,
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Normal'],
        fontSize=14,
        textColor=RGBColor(60, 60, 60),
        spaceAfter=12,
        alignment=TA_CENTER,
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=RGBColor(0, 0, 0),
        spaceAfter=6,
        spaceBefore=12,
    )
    
    # Title
    elements.append(Paragraph("Sarah Mitchell", title_style))
    elements.append(Paragraph("Senior Full-Stack Developer", subtitle_style))
    elements.append(Paragraph(
        "Email: sarah.mitchell@techmail.com | Phone: +1 (555) 987-6543 | San Francisco, CA",
        styles['Normal']
    ))
    elements.append(Spacer(1, 0.2*inch))
    
    # Professional Summary
    elements.append(Paragraph("PROFESSIONAL SUMMARY", heading_style))
    elements.append(Paragraph(
        "Highly skilled Full-Stack Developer with 5+ years of experience building scalable "
        "web applications and microservices. Expert in modern JavaScript frameworks, cloud "
        "architecture, and agile methodologies.",
        styles['Normal']
    ))
    elements.append(Spacer(1, 0.1*inch))
    
    # Technical Skills
    elements.append(Paragraph("TECHNICAL SKILLS", heading_style))
    skills = [
        "<b>Languages:</b> JavaScript, TypeScript, Python, Java, SQL, HTML5, CSS3",
        "<b>Frontend:</b> React, Next.js, Vue.js, Angular, Redux, Tailwind CSS",
        "<b>Backend:</b> Node.js, Express, Django, Flask, Spring Boot, FastAPI",
        "<b>Databases:</b> PostgreSQL, MongoDB, MySQL, Redis, DynamoDB",
        "<b>Cloud & DevOps:</b> AWS, Docker, Kubernetes, Jenkins, CI/CD, Terraform",
        "<b>Testing:</b> Jest, PyTest, Selenium, Cypress"
    ]
    for skill in skills:
        elements.append(Paragraph(skill, styles['Normal']))
    elements.append(Spacer(1, 0.1*inch))
    
    # Professional Experience
    elements.append(Paragraph("PROFESSIONAL EXPERIENCE", heading_style))
    elements.append(Paragraph(
        "<b>Senior Full-Stack Developer | TechCorp Solutions</b><br/>"
        "<i>San Francisco, CA | January 2022 - Present</i>",
        styles['Normal']
    ))
    
    responsibilities = [
        "Led development team of 8 engineers building enterprise SaaS platform",
        "Architected microservices infrastructure using Node.js, PostgreSQL, and AWS",
        "Reduced application load time by 60% through optimization",
        "Implemented CI/CD pipelines reducing deployment time by 75%"
    ]
    for resp in responsibilities:
        elements.append(Paragraph(f"• {resp}", styles['Normal']))
    
    elements.append(Spacer(1, 0.1*inch))
    
    # Education
    elements.append(Paragraph("EDUCATION", heading_style))
    elements.append(Paragraph(
        "<b>Bachelor of Science in Computer Science</b><br/>"
        "University of California, Berkeley | 2014 - 2018 | GPA: 3.7/4.0",
        styles['Normal']
    ))
    
    # Build PDF
    doc.build(elements)
    print(f"✅ Created: {output_path}")
    return output_path


if __name__ == "__main__":
    print("\n🔨 Generating Test Resume Files...")
    print("="*60)
    
    try:
        # Create DOCX files
        create_docx_resume_fullstack()
        create_docx_resume_devops()
        
        # Create PDF if reportlab is available
        if PDF_AVAILABLE:
            create_pdf_resume_fullstack()
        else:
            print("\n⚠️  To generate PDF files, install:")
            print("    pip install reportlab")
        
        print("\n="*60)
        print("✅ Test resume generation complete!")
        print("\nGenerated files in test-resumes/:")
        print("  • sarah_mitchell_fullstack.txt")
        print("  • sarah_mitchell_fullstack.docx")
        if PDF_AVAILABLE:
            print("  • sarah_mitchell_fullstack.pdf")
        print("  • michael_chen_devops.txt")
        print("  • michael_chen_devops.docx")
        
    except Exception as e:
        print(f"\n❌ Error generating files: {e}")
        import traceback
        traceback.print_exc()
